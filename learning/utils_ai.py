import pandas as pd
import os
import re
import sys
import traceback
import io
import json

# === 依赖库检测 ===
try:
    import docx
except ImportError:
    print("【提示】缺少 python-docx，请运行: pip install python-docx")

try:
    import pdfplumber
except ImportError:
    print("【提示】缺少 pdfplumber，请运行: pip install pdfplumber")

try:
    import dashscope
    from dashscope import Generation

    # TODO: 请务必在这里填入您在阿里云申请的 API-KEY (千万不要发到网上)
    dashscope.api_key = "sk-0ddb7aa22f694638ab3c9e4386e7f1b3"
except ImportError:
    print("【提示】缺少 dashscope，请运行: pip install dashscope")

from django.utils import timezone
from django.db import transaction
from learning.models import Exercise, Choice, KnowledgePoint, QMatrix


# ==========================================
# 1. 核心解析函数 (原版 - 用于处理排版好的Word)
# ==========================================
def parse_text_to_exercises(text):
    print(f"--- 开始解析文本，长度: {len(text)} ---")
    text = text.replace('\t', ' ')
    lines = text.split('\n')
    exercises = []

    current_exercise = None
    current_options = []

    pattern_question = re.compile(r'^(\d+[\.、\)]|【\d+】|Question\s*\d+)(.*)')
    pattern_option = re.compile(r'^(\(?\s*[A-D]\s*[\.、\)\s]?)(.*)', re.IGNORECASE)
    pattern_answer = re.compile(r'^(答案|Answer)[:：]\s*(.*)', re.IGNORECASE)
    pattern_kp = re.compile(r'^(知识点|考点|Knowledge)[:：]\s*(.*)', re.IGNORECASE)
    pattern_solution = re.compile(r'^(解析|Solution)[:：]\s*(.*)', re.IGNORECASE)

    for line in lines:
        line = line.strip()
        if not line: continue
        if line.lower() == 'undefined': continue

        q_match = pattern_question.match(line)
        if q_match:
            if current_exercise:
                current_exercise['options'] = current_options
                if not current_exercise['answer']: current_exercise['answer'] = '略'
                exercises.append(current_exercise)

            current_exercise = {
                'title': q_match.group(2).strip(),
                'answer': '',
                'knowledge_point': '',
                'solution': '',
                'options': []
            }
            current_options = []
            continue

        opt_match = pattern_option.match(line)
        if current_exercise and opt_match:
            raw_label = opt_match.group(1)
            label = re.sub(r'[\(\)\.、\s\[\]]', '', raw_label).upper()
            content = opt_match.group(2).strip()

            if label in ['A', 'B', 'C', 'D'] and content:
                current_options.append({'label': label, 'content': content})
            continue

        ans_match = pattern_answer.match(line)
        if current_exercise and ans_match:
            current_exercise['answer'] = ans_match.group(2).strip().upper()
            continue

        kp_match = pattern_kp.match(line)
        if current_exercise and kp_match:
            current_exercise['knowledge_point'] = kp_match.group(2).strip()
            continue

        sol_match = pattern_solution.match(line)
        if current_exercise and sol_match:
            current_exercise['solution'] = sol_match.group(2).strip()
            continue

    if current_exercise:
        current_exercise['options'] = current_options
        if not current_exercise['answer']: current_exercise['answer'] = '略'
        exercises.append(current_exercise)

    print(f"--- 解析完成，共识别出 {len(exercises)} 道题 ---")
    return exercises


# ==========================================
# 2. 文档处理入口 (直接导入Word/TXT模式)
# ==========================================
def handle_document_file(file_record):
    print(">>> 进入传统正则解析: handle_document_file")
    subject = file_record.subject
    teacher = file_record.teacher
    filename = file_record.original_filename.lower()

    full_text = ""
    count = 0

    try:
        try:
            with file_record.file.open('rb') as f:
                file_content = f.read()
        except Exception:
            with open(file_record.file.path, 'rb') as f:
                file_content = f.read()

        if not file_content: return 0

        if filename.endswith('.pdf'):
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    full_text = "\n".join([page.extract_text() or '' for page in pdf.pages])
            except Exception as e:
                print(f"PDF 解析失败: {e}")
        elif filename.endswith('.docx'):
            try:
                docx_file = io.BytesIO(file_content)
                doc = docx.Document(docx_file)
                full_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                print(f"Word解析异常: {e}")
                try:
                    full_text = file_content.decode('utf-8')
                except:
                    full_text = file_content.decode('gbk', errors='ignore')
        else:
            try:
                full_text = file_content.decode('utf-8')
            except:
                try:
                    full_text = file_content.decode('gbk')
                except:
                    return 0

        if not full_text.strip(): return 0

        exercises_data = parse_text_to_exercises(full_text)

        with transaction.atomic():
            for item in exercises_data:
                try:
                    q_type = 'single'
                    if not item['options']:
                        q_type = 'subjective'
                    elif len(item['answer']) > 1:
                        q_type = 'multiple'

                    opt_dict = {opt['label']: opt['content'] for opt in item['options']}
                    full_option_text = repr(opt_dict) if opt_dict else ""

                    exercise = Exercise.objects.create(
                        subject=subject,
                        title=item['title'][:50],
                        content=item['title'].replace('undefined', '').strip(),
                        question_type=q_type,
                        creator=teacher,
                        option_text=full_option_text,
                        answer=item['answer'],
                        solution=item.get('solution', ''),
                        problemsets="文件导入",
                        created_at=timezone.now()
                    )

                    for idx, opt in enumerate(item['options']):
                        labeled_content = f"{opt['label']}. {opt['content']}"
                        Choice.objects.create(
                            exercise=exercise,
                            content=labeled_content,
                            is_correct=(opt['label'] in item['answer']),
                            order=idx + 1
                        )

                    kp_str = item.get('knowledge_point', '')
                    if not kp_str: kp_str = "导入知识点"
                    kp_list = re.split(r'[，,、\s]+', kp_str)
                    for kp_name in kp_list:
                        if not kp_name: continue
                        kp, _ = KnowledgePoint.objects.get_or_create(subject=subject, name=kp_name,
                                                                     defaults={'parent': None})
                        QMatrix.objects.get_or_create(exercise=exercise, knowledge_point=kp, defaults={'weight': 1.0})
                    count += 1
                except Exception as e:
                    print(f"保存单条习题出错: {e}")
                    continue
    except Exception as e:
        print(f"【处理报错】: {e}")
        traceback.print_exc()
    finally:
        try:
            file_record.file.close()
        except:
            pass
    return count


# ==========================================
# 3. Excel/CSV 处理入口 - 维持原样
# ==========================================
def handle_data_file(file_record):
    print(">>> 进入表格解析: handle_data_file")
    count = 0
    file_path = file_record.file.path
    subject = file_record.subject
    teacher = file_record.teacher

    try:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path, encoding='utf-8-sig')
        else:
            df = pd.read_excel(file_path)

        with transaction.atomic():
            for _, row in df.iterrows():
                title = str(row.get('题目') or row.iloc[0]).strip()
                if len(title) < 2 or title == 'nan': continue

                def get_val(key, idx):
                    val = row.get(key)
                    if pd.isna(val): val = row.iloc[idx] if idx < len(row) else ''
                    return str(val).strip()

                opt_dict = {}
                for label, idx in [('A', 1), ('B', 2), ('C', 3), ('D', 4)]:
                    txt = get_val(f'选项{label}', idx)
                    if txt and txt != 'nan': opt_dict[label] = txt

                ans = str(row.get('答案') or 'A').strip().upper()

                exercise = Exercise.objects.create(
                    subject=subject,
                    title=title[:50],
                    content=title.replace('undefined', '').strip(),
                    question_type='single' if len(ans) == 1 else 'multiple',
                    creator=teacher,
                    option_text=repr(opt_dict),
                    answer=ans,
                    problemsets="Excel导入",
                    created_at=timezone.now()
                )

                for idx, (label, txt) in enumerate(opt_dict.items()):
                    Choice.objects.create(
                        exercise=exercise,
                        content=f"{label}. {txt}",
                        is_correct=(label in ans),
                        order=idx + 1
                    )
                count += 1
    except Exception as e:
        print(f"【Excel处理报错】: {e}")
        traceback.print_exc()
    return count


# ==========================================
# 4. AI 基于资料自动生成习题与知识点 (新增)
# ==========================================
def auto_generate_exercises_from_text(text):
    """
    调用通义千问，让它阅读资料并自动出题
    """
    print(">>> 开始调用通义千问基于资料自动生成习题...")

    # 截取前 4000 字（大模型阅读材料的范围，防止 Token 超限）
    content = text[:4000]

    # 专门针对计算机科学类的出题 Prompt
    prompt = f"""
    你是一个专业的计算机科学大学教师。请阅读以下教学资料，并根据资料中的核心内容，生成 3 道高质量的单选题。
    要求：
    1. 题目要有针对性，选项要有迷惑性。
    2. 为每道题提取 1-2 个核心知识点短语。
    3. 必须且只能返回一个纯 JSON 数组格式，不要包含任何前缀、后缀或 Markdown 代码块标记（如```json）。

    返回的 JSON 必须严格遵守以下格式示例：
    [
        {{
            "title": "MVC设计模式中，负责处理业务逻辑和数据库交互的是哪一部分？",
            "options": {{
                "A": "Model (模型)",
                "B": "View (视图)",
                "C": "Controller (控制器)",
                "D": "Template (模板)"
            }},
            "answer": "A",
            "solution": "在MVC架构中，Model负责数据的封装和业务逻辑处理。",
            "knowledge_points": ["MVC架构", "后端开发基础"]
        }}
    ]

    请根据以下资料内容出题：
    {content}
    """

    try:
        response = Generation.call(
            model="qwen-turbo",
            prompt=prompt,
            result_format='message'
        )

        if response.status_code == 200:
            result_text = response.output.choices[0].message.content.strip()

            # 清洗并解析 JSON
            clean_json = result_text.replace('```json', '').replace('```', '').strip()
            exercise_list = json.loads(clean_json)

            return exercise_list if isinstance(exercise_list, list) else []
        else:
            print(f"大模型调用失败: {response.code} - {response.message}")
            return []
    except Exception as e:
        print(f"AI 生成习题解析异常: {e}")
        return []


# ==========================================
# 5. 资料处理入口 (提取文字 -> AI出题 -> 存入题库)
# ==========================================
def handle_material_generation(file_record):
    """
    读取上传的资料，调用 AI 出题，并将题目和知识点无缝存入现有的数据库
    """
    print(">>> 进入大模型出题: handle_material_generation 函数")
    subject = file_record.subject
    teacher = getattr(file_record, 'teacher', getattr(file_record, 'uploader', None))
    filename = file_record.original_filename.lower()
    full_text = ""

    # 1. 提取资料文本
    try:
        try:
            with file_record.file.open('rb') as f:
                file_content = f.read()
        except Exception:
            with open(file_record.file.path, 'rb') as f:
                file_content = f.read()

        if not file_content: return 0

        if filename.endswith('.pdf'):
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    full_text = "\n".join([page.extract_text() or '' for page in pdf.pages])
            except Exception as e:
                print(f"资料 PDF 解析失败: {e}")
        elif filename.endswith('.docx'):
            try:
                docx_file = io.BytesIO(file_content)
                doc = docx.Document(docx_file)
                full_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                print(f"资料 Word 解析异常: {e}")
                try:
                    full_text = file_content.decode('utf-8')
                except:
                    full_text = file_content.decode('gbk', errors='ignore')
        else:
            try:
                full_text = file_content.decode('utf-8')
            except:
                try:
                    full_text = file_content.decode('gbk')
                except:
                    return 0

        if not full_text.strip(): return 0

        # 2. 召唤 AI 自动出题
        generated_exercises = auto_generate_exercises_from_text(full_text)
        print(f"--- AI 成功生成了 {len(generated_exercises)} 道题 ---")

        count = 0
        with transaction.atomic():
            for item in generated_exercises:
                try:
                    opt_dict = item.get('options', {})
                    full_option_text = repr(opt_dict) if opt_dict else ""

                    exercise = Exercise.objects.create(
                        subject=subject,
                        title=item.get('title', '')[:200],
                        content=item.get('title', ''),
                        question_type='single',
                        creator_id=teacher.id if teacher else 1,
                        option_text=full_option_text,
                        answer=item.get('answer', 'A'),
                        solution=item.get('solution', ''),
                        problemsets="AI智能生成",
                        created_at=timezone.now()
                    )

                    for idx, (label, content) in enumerate(opt_dict.items()):
                        Choice.objects.create(
                            exercise=exercise,
                            content=f"{label}. {content}",
                            is_correct=(label == item.get('answer')),
                            order=idx + 1
                        )

                    kps = item.get('knowledge_points', [])
                    for kp_name in kps:
                        if not kp_name.strip(): continue
                        kp, _ = KnowledgePoint.objects.get_or_create(
                            subject=subject,
                            name=kp_name.strip(),
                            defaults={'parent': None}
                        )
                        QMatrix.objects.get_or_create(exercise=exercise, knowledge_point=kp, defaults={'weight': 1.0})

                    count += 1
                except Exception as e:
                    print(f"保存单条 AI 生成习题出错: {e}")
                    continue

        return count

    except Exception as e:
        print(f"处理资料出题出错: {e}")
        traceback.print_exc()
        return 0


# ==========================================
# 6. 🎯 终极管家接口 (根据情况自动分配)
# ==========================================
def smart_handle_upload(file_record, mode='ai'):
    """
    自动判断是直接提取还是让AI出题。
    如果想要强制用AI出题，mode保持 'ai' 即可。
    """
    filename = file_record.original_filename.lower()

    if mode == 'ai':
        return handle_material_generation(file_record)
    else:
        if filename.endswith(('.xls', '.xlsx', '.csv')):
            return handle_data_file(file_record)
        else:
            return handle_document_file(file_record)



# ==========================================
# 7. PDF课本完整分析功能
# ==========================================
def analyze_textbook_pdf(textbook_builder):
    """
    分析PDF课本，提取章节结构、知识点、习题，并构建知识图谱
    """
    print(f">>> 开始分析PDF课本: {textbook_builder.subject_name}")
    
    try:
        # 1. 提取PDF文本
        textbook_builder.status = 'extracting_text'
        textbook_builder.save()
        
        file_path = textbook_builder.textbook_file.path
        full_text = ""
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                full_text += text + "\n"
        
        textbook_builder.extracted_text = full_text[:10000]  # 保存前10000字符用于参考
        textbook_builder.status = 'analyzing_content'
        textbook_builder.save()
        
        # 2. 分析章节结构
        chapters = extract_chapters_from_text(full_text)
        textbook_builder.chapter_count = len(chapters)
        textbook_builder.status = 'generating_exercises'
        textbook_builder.save()
        
        # 3. 创建课程科目
        subject, created = Subject.objects.get_or_create(
            name=textbook_builder.subject_name,
            defaults={
                'description': textbook_builder.subject_description or f"基于《{textbook_builder.subject_name}》PDF课本快速构建的课程"
            }
        )
        
        # 4. 教师关联课程
        TeacherSubject.objects.get_or_create(
            teacher=textbook_builder.teacher,
            subject=subject
        )
        
        textbook_builder.generated_subject = subject
        textbook_builder.status = 'extracting_knowledge'
        textbook_builder.save()
        
        # 5. 提取知识点
        knowledge_points = extract_knowledge_points_from_text(full_text, subject)
        textbook_builder.knowledge_point_count = len(knowledge_points)
        textbook_builder.status = 'building_graph'
        textbook_builder.save()
        
        # 6. 构建知识点关系
        relationships = build_knowledge_graph(knowledge_points, subject)
        textbook_builder.relationship_count = len(relationships)
        
        # 7. 生成习题
        exercises = generate_exercises_from_textbook(full_text, subject, textbook_builder.teacher)
        textbook_builder.exercise_count = len(exercises)
        
        # 8. 完成初步处理，设置为待审核状态
        textbook_builder.status = 'review_pending'
        textbook_builder.processed_at = timezone.now()
        textbook_builder.save()
        
        return {
            'success': True,
            'subject_id': subject.id,
            'chapters': len(chapters),
            'knowledge_points': len(knowledge_points),
            'relationships': len(relationships),
            'exercises': len(exercises)
        }
        
    except Exception as e:
        print(f"PDF课本分析失败: {e}")
        traceback.print_exc()
        textbook_builder.status = 'error'
        textbook_builder.error_message = str(e)
        textbook_builder.save()
        return {'success': False, 'error': str(e)}


def extract_chapters_from_text(text):
    """从文本中提取章节结构"""
    chapters = []
    
    # 匹配章节标题模式
    chapter_patterns = [
        r'第[一二三四五六七八九十\d]+章\s+[^\n]+',
        r'Chapter\s+\d+\s*[:：]?\s*[^\n]+',
        r'\d+\.\d*\s+[^\n]+',  # 1.1 标题
        r'[一二三四五六七八九十]+、\s*[^\n]+',
    ]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        if len(line) > 100:  # 跳过过长的行（可能是段落）
            continue
            
        for pattern in chapter_patterns:
            if re.match(pattern, line):
                chapters.append({
                    'title': line,
                    'content_start': i
                })
                break
    
    return chapters


def extract_knowledge_points_from_text(text, subject):
    """从文本中提取知识点"""
    print(">>> 开始提取知识点...")
    
    # 使用AI提取知识点
    prompt = f"""
    你是一个专业的课程内容分析师。请从以下教材内容中提取核心知识点。
    
    要求：
    1. 提取20-30个最重要的知识点
    2. 每个知识点用简洁的短语表示
    3. 知识点之间要有层次关系（父子关系）
    4. 返回JSON格式
    
    教材内容（摘要）：
    {text[:3000]}
    
    请返回以下格式的JSON：
    {{
        "knowledge_points": [
            {{
                "name": "知识点名称",
                "parent": "父知识点名称（可为空）",
                "description": "知识点简要描述"
            }}
        ]
    }}
    """
    
    try:
        response = Generation.call(
            model="qwen-turbo",
            prompt=prompt,
            result_format='message'
        )
        
        if response.status_code == 200:
            result_text = response.output.choices[0].message.content.strip()
            clean_json = result_text.replace('```json', '').replace('```', '').strip()
            data = json.loads(clean_json)
            
            knowledge_points = []
            kp_map = {}  # 名称到对象的映射
            
            for kp_data in data.get('knowledge_points', []):
                # 创建知识点
                parent = None
                if kp_data.get('parent'):
                    # 先检查父知识点是否已创建
                    if kp_data['parent'] in kp_map:
                        parent = kp_map[kp_data['parent']]
                    else:
                        # 如果父知识点不存在，先创建它
                        parent_kp, _ = KnowledgePoint.objects.get_or_create(
                            subject=subject,
                            name=kp_data['parent'],
                            defaults={'parent': None}
                        )
                        kp_map[kp_data['parent']] = parent_kp
                        parent = parent_kp
                
                kp, created = KnowledgePoint.objects.get_or_create(
                    subject=subject,
                    name=kp_data['name'],
                    defaults={
                        'parent': parent,
                        'description': kp_data.get('description', '')
                    }
                )
                
                kp_map[kp_data['name']] = kp
                knowledge_points.append(kp)
            
            return knowledge_points
            
    except Exception as e:
        print(f"知识点提取失败: {e}")
    
    # 如果AI提取失败，使用简单的关键词提取
    return extract_keywords_as_knowledge_points(text, subject)


def extract_keywords_as_knowledge_points(text, subject):
    """使用关键词提取作为备选方案"""
    # 简单的关键词提取（实际项目中可以使用更复杂的NLP技术）
    keywords = re.findall(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*', text[:5000])
    
    knowledge_points = []
    for i, keyword in enumerate(keywords[:20]):  # 取前20个作为知识点
        if len(keyword) > 3 and len(keyword) < 50:
            kp, _ = KnowledgePoint.objects.get_or_create(
                subject=subject,
                name=keyword,
                defaults={'parent': None}
            )
            knowledge_points.append(kp)
    
    return knowledge_points


def build_knowledge_graph(knowledge_points, subject):
    """构建知识点关系图"""
    print(">>> 开始构建知识点关系图...")
    
    relationships = []
    
    # 简单的基于文本相似度的关系构建
    # 实际项目中可以使用更复杂的NLP技术
    for i, kp1 in enumerate(knowledge_points):
        for j, kp2 in enumerate(knowledge_points):
            if i >= j:  # 避免重复和自环
                continue
                
            # 简单的名称相似度检查
            words1 = set(kp1.name.lower().split())
            words2 = set(kp2.name.lower().split())
            common_words = words1.intersection(words2)
            
            if common_words and len(common_words) / min(len(words1), len(words2)) > 0.3:
                # 创建关系
                relationship, created = KnowledgeGraph.objects.get_or_create(
                    subject=subject,
                    source=kp1,
                    target=kp2
                )
                relationships.append(relationship)
    
    return relationships


def generate_exercises_from_textbook(text, subject, teacher):
    """从课本内容生成习题"""
    print(">>> 开始从课本内容生成习题...")
    
    # 使用AI生成习题
    prompt = f"""
    你是一个专业的教师。请根据以下教材内容，生成10道高质量的习题。
    
    要求：
    1. 包含单选题、多选题和判断题
    2. 题目要覆盖教材的核心内容
    3. 每道题都要有明确的答案和解析
    4. 每道题关联1-2个知识点
    5. 返回JSON格式
    
    教材内容（摘要）：
    {text[:4000]}
    
    请返回以下格式的JSON：
    {{
        "exercises": [
            {{
                "title": "习题标题",
                "content": "习题内容",
                "question_type": "single/multiple/judgment",
                "options": {{
                    "A": "选项A内容",
                    "B": "选项B内容",
                    "C": "选项C内容",
                    "D": "选项D内容"
                }},
                "answer": "A",  # 多选题用"AB"格式
                "solution": "答案解析",
                "knowledge_points": ["知识点1", "知识点2"]
            }}
        ]
    }}
    """
    
    try:
        response = Generation.call(
            model="qwen-turbo",
            prompt=prompt,
            result_format='message'
        )
        
        if response.status_code == 200:
            result_text = response.output.choices[0].message.content.strip()
            clean_json = result_text.replace('```json', '').replace('```', '').strip()
            data = json.loads(clean_json)
            
            exercises = []
            
            with transaction.atomic():
                for ex_data in data.get('exercises', []):
                    try:
                        opt_dict = ex_data.get('options', {})
                        full_option_text = repr(opt_dict) if opt_dict else ""
                        
                        # 确定题型
                        q_type = ex_data.get('question_type', 'single')
                        if q_type == 'judgment':
                            q_type = 'judgment'
                            opt_dict = {'A': '正确', 'B': '错误'} if not opt_dict else opt_dict
                        
                        exercise = Exercise.objects.create(
                            subject=subject,
                            title=ex_data.get('title', '')[:200],
                            content=ex_data.get('content', ''),
                            question_type=q_type,
                            creator=teacher,
                            option_text=full_option_text,
                            answer=ex_data.get('answer', 'A'),
                            solution=ex_data.get('solution', ''),
                            problemsets="课本快速构建",
                            created_at=timezone.now()
                        )
                        
                        # 创建选项
                        for idx, (label, content) in enumerate(opt_dict.items()):
                            is_correct = False
                            if q_type == 'judgment':
                                is_correct = (label == 'A' and ex_data.get('answer') == '正确') or \
                                           (label == 'B' and ex_data.get('answer') == '错误')
                            else:
                                is_correct = label in ex_data.get('answer', '')
                                
                            Choice.objects.create(
                                exercise=exercise,
                                content=f"{label}. {content}",
                                is_correct=is_correct,
                                order=idx + 1
                            )
                        
                        # 关联知识点
                        for kp_name in ex_data.get('knowledge_points', []):
                            if kp_name.strip():
                                kp, _ = KnowledgePoint.objects.get_or_create(
                                    subject=subject,
                                    name=kp_name.strip(),
                                    defaults={'parent': None}
                                )
                                QMatrix.objects.get_or_create(
                                    exercise=exercise,
                                    knowledge_point=kp,
                                    defaults={'weight': 1.0}
                                )
                        
                        exercises.append(exercise)
                        
                    except Exception as e:
                        print(f"保存习题失败: {e}")
                        continue
            
            return exercises
            
    except Exception as e:
        print(f"习题生成失败: {e}")
    
    return []


# ==========================================
# 8. 课本快速构建接口
# ==========================================
def quick_build_course_from_textbook(textbook_builder):
    """
    快速构建课程的入口函数
    """
    return analyze_textbook_pdf(textbook_builder)